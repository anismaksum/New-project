from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "PRD_SmartWaste_Muhammad_Anis_Maksum_Winarso.docx"

GREEN_DARK = "0B3B2E"
GREEN = "0E7A50"
GREEN_SOFT = "EAF3EA"
LIME = "A9D95A"
AMBER = "F2A93B"
CORAL = "E86D5A"
INK = "10231D"
MUTED = "66756F"
BORDER = "DDE8DE"
BG = "F6F8F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_cm: float) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 24, GREEN_DARK),
        ("Subtitle", 11, MUTED),
        ("Heading 1", 16, GREEN_DARK),
        ("Heading 2", 13, GREEN),
        ("Heading 3", 11, INK),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Product Requirement Document - SmartWaste"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.runs[0].font.name = "Arial"
    header_para.runs[0].font.size = Pt(8)
    header_para.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BORDER)
    border.append(bottom)
    header_para._p.get_or_add_pPr().append(border)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, "List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        add_paragraph(document, item, "List Number")


def add_callout(document: Document, title: str, body: str, fill: str = GREEN_SOFT) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, 17)
    set_table_borders(table, fill)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_run = body_paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(9.8)
    body_run.font.color.rgb = RGBColor.from_string(INK)


def add_simple_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width(table, sum(widths_cm))
    set_table_borders(table)
    table.rows[0].height_rule = None
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_width(cell, widths_cm[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, GREEN_SOFT)
        set_cell_text(cell, header, bold=True, color=GREEN_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, widths_cm[idx])
            set_cell_margins(cell)
            set_cell_text(cell, text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    document.add_paragraph()


def add_cover(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Product Requirement Document")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(GREEN_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Aplikasi SmartWaste")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(GREEN)

    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Panduan fitur, fungsi, tujuan, perilaku aplikasi, dan cara pengumpulan")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph()
    add_callout(
        document,
        "Ringkasan dokumen",
        "Dokumen ini menjelaskan kebutuhan produk SmartWaste versi mobile: aplikasi pengelolaan setoran sampah bernilai, reward Eco Pts, kalkulator poin, penjadwalan pickup, panduan sortasi, serta profil pengguna. PRD ini juga memuat cara pengumpulan kebutuhan dan tata cara pengumpulan laporan/proyek.",
    )
    document.add_paragraph()

    metadata_rows = [
        ["Nama Produk", "SmartWaste"],
        ["Jenis Dokumen", "Product Requirement Document (PRD)"],
        ["Versi", "1.0"],
        ["Tanggal", date.today().strftime("%d %B %Y")],
        ["Penyusun", "Muhammad Anis Maksum Winarso"],
        ["Platform", "Flutter Android, Web preview, dan cross-platform ready"],
        ["Output Build", "Android App Bundle (.aab) release"],
    ]
    add_simple_table(document, ["Informasi", "Detail"], metadata_rows, [5, 12])
    document.add_page_break()


def add_table_of_contents(document: Document) -> None:
    add_paragraph(document, "Daftar Isi", "Heading 1")
    add_numbered(
        document,
        [
            "Ringkasan Eksekutif",
            "Latar Belakang dan Masalah",
            "Tujuan Produk dan Sasaran",
            "Ruang Lingkup Produk",
            "Target Pengguna",
            "Fitur dan Perilaku Aplikasi",
            "Kebutuhan Fungsional",
            "Kebutuhan Non-Fungsional",
            "Alur Pengguna dan Aturan Bisnis",
            "Spesifikasi UI/UX",
            "Pengujian dan Kriteria Penerimaan",
            "Cara Pengumpulan Kebutuhan",
            "Cara Pengumpulan/Penyerahan Tugas",
            "Lampiran Teknis",
        ],
    )
    document.add_page_break()


def build_report() -> Path:
    DOCS_DIR.mkdir(exist_ok=True)
    document = Document()
    configure_document(document)
    add_cover(document)
    add_table_of_contents(document)

    add_paragraph(document, "1. Ringkasan Eksekutif", "Heading 1")
    add_paragraph(
        document,
        "SmartWaste adalah aplikasi mobile berbasis Flutter untuk membantu pengguna menyetorkan sampah bernilai secara lebih terarah. Aplikasi menggabungkan informasi kategori sampah, estimasi reward Eco Pts, penjadwalan penjemputan kurir, panduan sortasi, serta profil pengguna dalam satu pengalaman yang ringan dan mudah dipahami.",
    )
    add_callout(
        document,
        "Visi produk",
        "Membuat aktivitas memilah dan menyetor sampah terasa sederhana, terukur, dan memberi umpan balik yang jelas bagi pengguna.",
    )

    add_paragraph(document, "2. Latar Belakang dan Masalah", "Heading 1")
    add_paragraph(
        document,
        "Pengelolaan sampah rumah tangga sering terkendala oleh minimnya informasi kategori, ketidakjelasan nilai tukar, dan proses penjemputan yang belum praktis. Pengguna membutuhkan aplikasi yang tidak hanya tampil menarik, tetapi juga membantu mengambil keputusan: sampah apa yang disetor, berapa estimasi poinnya, dan kapan kurir dapat mengambilnya.",
    )
    add_bullets(
        document,
        [
            "Pengguna belum selalu mengetahui jenis sampah yang dapat didaur ulang.",
            "Estimasi poin/reward sering tidak terlihat sebelum setoran dilakukan.",
            "Penjadwalan pickup manual dapat menyebabkan miskomunikasi waktu dan catatan lokasi.",
            "Informasi riwayat dan status pengguna belum tersaji ringkas.",
        ],
    )

    add_paragraph(document, "3. Tujuan Produk dan Sasaran", "Heading 1")
    add_simple_table(
        document,
        ["Tujuan", "Indikator Keberhasilan"],
        [
            ["Memudahkan pemilahan sampah", "Pengguna dapat memilih kategori sampah dan membaca tips sortasi dalam kurang dari 2 tap."],
            ["Menghitung reward secara transparan", "Pengguna dapat melihat estimasi Eco Pts berdasarkan kategori dan berat."],
            ["Mempercepat proses pickup", "Pengguna dapat memilih slot jemput dan menambahkan catatan kurir."],
            ["Meningkatkan motivasi pengguna", "Dashboard menampilkan total Eco Pts, progress reward, dan ringkasan pickup."],
        ],
        [6, 11],
    )

    add_paragraph(document, "4. Ruang Lingkup Produk", "Heading 1")
    add_paragraph(document, "Termasuk dalam versi 1.0", "Heading 2")
    add_bullets(
        document,
        [
            "Login demo dengan validasi input.",
            "Dashboard reward dan metrik ringkas.",
            "Kategori sampah: plastik, logam, kertas, kaca.",
            "Kalkulator Eco Pts berdasarkan berat sampah.",
            "Catat setoran untuk menambah total Eco Pts.",
            "Penjadwalan pickup dengan slot waktu dan catatan.",
            "Panduan sortasi per kategori.",
            "Profil pengguna, reminder pickup, status pickup terakhir, dan logout.",
        ],
    )
    add_paragraph(document, "Di luar lingkup versi 1.0", "Heading 2")
    add_bullets(
        document,
        [
            "Autentikasi backend real-time.",
            "Peta lokasi GPS dan pelacakan kurir langsung.",
            "Pembayaran atau penukaran voucher sungguhan.",
            "Integrasi database cloud dan notifikasi push asli.",
        ],
    )

    add_paragraph(document, "5. Target Pengguna", "Heading 1")
    add_simple_table(
        document,
        ["Persona", "Kebutuhan", "Perilaku yang Didukung"],
        [
            ["Warga rumah tangga", "Menyetor sampah daur ulang dengan cepat.", "Memilih kategori, menghitung poin, menjadwalkan pickup."],
            ["Petugas bank sampah", "Melihat kategori dan catatan setoran.", "Membaca catatan pickup dan memahami jenis sampah yang disiapkan."],
            ["Mahasiswa/pengguna edukasi", "Memahami alur aplikasi smart waste.", "Mempelajari fitur, UI/UX, dan proses build aplikasi."],
        ],
        [4, 6, 7],
    )

    add_paragraph(document, "6. Fitur dan Perilaku Aplikasi", "Heading 1")
    add_paragraph(document, "6.1 Login", "Heading 2")
    add_bullets(
        document,
        [
            "Pengguna memasukkan username dan password.",
            "Aplikasi menyediakan tombol Isi Demo untuk mengisi akun contoh admin/12345.",
            "Password dapat ditampilkan atau disembunyikan.",
            "Jika kredensial benar, pengguna masuk ke dashboard; jika salah, aplikasi menampilkan snackbar informasi.",
        ],
    )
    add_paragraph(document, "6.2 Dashboard Beranda", "Heading 2")
    add_bullets(
        document,
        [
            "Menampilkan sapaan, total Eco Pts, progress menuju voucher, jumlah pickup, dan estimasi penghematan CO2.",
            "Menyediakan shortcut Scan, Bank Sampah, Riwayat, dan Jemput.",
            "Menampilkan kartu kategori sampah dengan gambar lokal agar tetap tampil tanpa jaringan internet.",
        ],
    )
    add_paragraph(document, "6.3 Smart Tools", "Heading 2")
    add_bullets(
        document,
        [
            "Pengguna memilih kategori sampah melalui choice chip.",
            "Pengguna mengatur berat sampah memakai slider atau tombol tambah/kurang.",
            "Aplikasi menghitung Eco Pts berdasarkan rumus: berat kg x poin per kg.",
            "Tombol Catat Setoran menambahkan poin ke total pengguna.",
            "Pengguna dapat memilih slot pickup dan menambahkan catatan untuk kurir.",
            "Switch reminder pickup dapat diaktifkan atau dinonaktifkan.",
        ],
    )
    add_paragraph(document, "6.4 Profil", "Heading 2")
    add_bullets(
        document,
        [
            "Menampilkan identitas pengguna, program studi, kampus, total Eco Pts, dan jumlah pickup.",
            "Menampilkan status reminder, jadwal pickup terakhir, dan catatan pickup.",
            "Logout mengembalikan pengguna ke halaman login dan membersihkan stack navigasi.",
        ],
    )

    add_paragraph(document, "7. Kebutuhan Fungsional", "Heading 1")
    add_simple_table(
        document,
        ["ID", "Kebutuhan", "Prioritas", "Kriteria Penerimaan"],
        [
            ["FR-01", "Aplikasi harus menyediakan login demo.", "Must", "Akun admin/12345 dapat masuk ke dashboard."],
            ["FR-02", "Aplikasi harus memvalidasi input login.", "Must", "Input kosong atau salah menampilkan pesan yang jelas."],
            ["FR-03", "Dashboard harus menampilkan total Eco Pts.", "Must", "Nilai poin terlihat di kartu reward utama."],
            ["FR-04", "Pengguna harus dapat memilih kategori sampah.", "Must", "Kategori terpilih terlihat secara visual."],
            ["FR-05", "Kalkulator harus menghitung Eco Pts.", "Must", "Perubahan berat mengubah estimasi poin."],
            ["FR-06", "Pengguna harus dapat mencatat setoran.", "Should", "Total Eco Pts bertambah dan snackbar muncul."],
            ["FR-07", "Pengguna harus dapat menjadwalkan pickup.", "Must", "Slot dan catatan tersimpan di profil."],
            ["FR-08", "Aplikasi harus menyediakan panduan sortasi.", "Should", "Tips tiap kategori dapat dibuka di halaman Tools."],
            ["FR-09", "Aplikasi harus menyediakan logout.", "Must", "Logout membawa pengguna ke halaman login."],
        ],
        [1.6, 5.2, 2.1, 8.1],
    )

    add_paragraph(document, "8. Kebutuhan Non-Fungsional", "Heading 1")
    add_simple_table(
        document,
        ["Aspek", "Kebutuhan"],
        [
            ["Kinerja", "Transisi halaman dan interaksi kalkulator harus terasa responsif pada perangkat Android umum."],
            ["Aksesibilitas", "Kontras teks harus cukup, tombol memiliki ikon yang familiar, dan label input mudah dipahami."],
            ["Keandalan", "Aset kategori memakai gambar lokal agar UI tidak rusak saat offline."],
            ["Maintainability", "Tema warna dan style dipusatkan di AppTheme agar konsisten dan mudah diubah."],
            ["Portabilitas", "Kode Flutter siap dijalankan pada Android dan web preview."],
            ["Keamanan", "Versi demo tidak menyimpan data sensitif; autentikasi produksi perlu backend dan hashing password."],
        ],
        [4, 13],
    )

    add_paragraph(document, "9. Alur Pengguna dan Aturan Bisnis", "Heading 1")
    add_paragraph(document, "Alur utama pengguna", "Heading 2")
    add_numbered(
        document,
        [
            "Pengguna membuka aplikasi dan masuk menggunakan akun demo.",
            "Pengguna melihat dashboard reward dan memilih kategori sampah.",
            "Pengguna membuka tab Tools untuk menghitung poin berdasarkan berat.",
            "Pengguna mencatat setoran sehingga total Eco Pts bertambah.",
            "Pengguna menjadwalkan pickup dan menambahkan catatan untuk kurir.",
            "Pengguna membuka profil untuk melihat ringkasan status dan melakukan logout.",
        ],
    )
    add_paragraph(document, "Aturan bisnis utama", "Heading 2")
    add_simple_table(
        document,
        ["Kategori", "Poin per Kg", "Catatan Sortasi"],
        [
            ["Plastik", "50", "Bilas, keringkan, dan pipihkan."],
            ["Logam", "120", "Pisahkan dari sampah basah dan material tajam."],
            ["Kertas", "30", "Ikat rapi dan jauhkan dari minyak/air."],
            ["Kaca", "80", "Bungkus pecahan dan beri tanda."],
        ],
        [4, 3, 10],
    )

    add_paragraph(document, "10. Spesifikasi UI/UX", "Heading 1")
    add_bullets(
        document,
        [
            "Gaya visual menggunakan eco-minimalism yang bersih, terang, dan modern.",
            "Warna utama hijau tua dan hijau aksen, dilengkapi amber, teal, dan coral agar kategori tidak monoton.",
            "Komponen utama memakai radius 8 px agar tampil rapi dan konsisten.",
            "Navigasi bawah terdiri dari Beranda, Tools, dan Profil untuk memisahkan tugas utama.",
            "Kontrol interaktif memakai pola yang familiar: slider untuk berat, switch untuk reminder, dropdown untuk slot pickup, dan chip untuk kategori.",
            "Konten penting diletakkan di atas: reward, metrik, dan shortcut tindakan cepat.",
        ],
    )
    add_callout(
        document,
        "Prinsip pengalaman",
        "Pengguna tidak perlu membaca instruksi panjang di layar. Setiap fitur memakai label singkat, ikon familiar, dan umpan balik langsung melalui snackbar, dialog, atau perubahan status.",
        "FFF6E5",
    )

    add_paragraph(document, "11. Pengujian dan Kriteria Penerimaan", "Heading 1")
    add_simple_table(
        document,
        ["Area", "Skenario", "Status"],
        [
            ["Static analysis", "Menjalankan flutter analyze.", "Lulus, no issues found."],
            ["Widget test", "Login demo membuka dashboard.", "Lulus."],
            ["Widget test", "Tools dapat mencatat setoran.", "Lulus."],
            ["Build release", "Menjalankan flutter build appbundle.", "Lulus, menghasilkan app-release.aab."],
            ["Preview web", "Server lokal merespons http://127.0.0.1:54545.", "Lulus, HTTP 200."],
        ],
        [3.5, 9, 4.5],
    )
    add_paragraph(document, "Kriteria penerimaan versi 1.0", "Heading 2")
    add_bullets(
        document,
        [
            "Aplikasi dapat dibuka dan pengguna dapat login memakai akun demo.",
            "Dashboard menampilkan reward, kategori, dan shortcut tanpa error gambar.",
            "Kalkulator poin menghasilkan nilai sesuai kategori dan berat.",
            "Catat setoran menambah total Eco Pts.",
            "Jadwal pickup tersimpan dan tampil pada profil.",
            "Build .aab selesai tanpa error Gradle.",
        ],
    )

    add_paragraph(document, "12. Cara Pengumpulan Kebutuhan", "Heading 1")
    add_paragraph(
        document,
        "Cara pengumpulan kebutuhan dipakai untuk memastikan fitur SmartWaste benar-benar sesuai dengan masalah pengguna dan proses operasional bank sampah.",
    )
    add_simple_table(
        document,
        ["Metode", "Tujuan", "Output"],
        [
            ["Observasi", "Melihat alur pemilahan, penimbangan, dan pickup sampah.", "Catatan proses, titik masalah, dan peluang fitur."],
            ["Wawancara", "Menggali kebutuhan warga, petugas, dan pengelola bank sampah.", "Daftar kebutuhan pengguna dan prioritas."],
            ["Studi dokumen", "Mempelajari daftar kategori, harga/poin, dan aturan sortasi.", "Aturan bisnis kategori sampah."],
            ["Benchmarking", "Membandingkan aplikasi atau layanan serupa.", "Referensi fitur dan standar UX."],
            ["Uji prototipe", "Memvalidasi apakah pengguna paham alur aplikasi.", "Masukan UI/UX dan daftar perbaikan."],
        ],
        [3.5, 7, 6.5],
    )

    add_paragraph(document, "13. Cara Pengumpulan/Penyerahan Tugas", "Heading 1")
    add_paragraph(
        document,
        "Bagian ini menjelaskan cara mengumpulkan hasil pekerjaan aplikasi SmartWaste agar pemeriksaan tugas lebih mudah dan lengkap.",
    )
    add_numbered(
        document,
        [
            "Pastikan folder proyek Flutter sudah dapat dibuka dan menjalankan flutter pub get.",
            "Pastikan hasil analisis bersih dengan perintah flutter analyze.",
            "Pastikan pengujian lulus dengan perintah flutter test.",
            "Buat Android App Bundle dengan perintah flutter build appbundle.",
            "Siapkan file PDF PRD sebagai laporan utama.",
            "Kumpulkan source code proyek dalam bentuk folder atau ZIP.",
            "Kumpulkan file app-release.aab dari folder build/app/outputs/bundle/release.",
            "Gunakan penamaan file yang jelas, misalnya PRD_SmartWaste_Muhammad_Anis_Maksum_Winarso.pdf dan smartwaste_source.zip.",
            "Upload ke LMS, Google Drive, atau media pengumpulan yang diminta dosen/guru.",
            "Jika memakai Google Drive, pastikan permission file dapat diakses oleh pemeriksa.",
        ],
    )
    add_paragraph(document, "Checklist pengumpulan", "Heading 2")
    add_simple_table(
        document,
        ["No", "Item", "Status"],
        [
            ["1", "PDF laporan PRD", "Wajib"],
            ["2", "Source code Flutter", "Wajib"],
            ["3", "File app-release.aab", "Wajib"],
            ["4", "Screenshot aplikasi atau link preview", "Opsional"],
            ["5", "Catatan akun demo: admin / 12345", "Disarankan"],
        ],
        [1.2, 11, 4.8],
    )

    add_paragraph(document, "14. Lampiran Teknis", "Heading 1")
    add_simple_table(
        document,
        ["Item", "Keterangan"],
        [
            ["Framework", "Flutter 3.41.4, Dart 3.11.1"],
            ["Entry point", "lib/main.dart"],
            ["Screen utama", "login_screen.dart, home_screen.dart, profile_screen.dart"],
            ["Tema", "lib/theme/app_theme.dart"],
            ["Aset lokal", "assets/images.jpg, download(2).jpg, download(3).jpg, wew.jpg"],
            ["Build AAB", "build/app/outputs/bundle/release/app-release.aab"],
            ["Akun demo", "Username: admin, Password: 12345"],
        ],
        [4.5, 12.5],
    )
    add_paragraph(document, "Rekomendasi pengembangan berikutnya", "Heading 2")
    add_bullets(
        document,
        [
            "Menambahkan backend untuk akun pengguna, riwayat setoran, dan data pickup real.",
            "Mengintegrasikan notifikasi push untuk pengingat pickup.",
            "Menambahkan peta lokasi bank sampah dan tracking kurir.",
            "Menambahkan fitur penukaran Eco Pts ke voucher atau saldo.",
            "Menambahkan dashboard admin untuk memvalidasi setoran dan mengelola kategori.",
        ],
    )

    document.core_properties.title = "PRD SmartWaste"
    document.core_properties.subject = "Product Requirement Document aplikasi SmartWaste"
    document.core_properties.author = "Muhammad Anis Maksum Winarso"
    document.core_properties.comments = "Generated for SmartWaste Flutter project"
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_report())
